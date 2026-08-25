# import neccessary functions, libraries, and packages
from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas
from django.contrib.auth.decorators import login_required


def _metric_line(label, before, after):
    b = f"{before:.3f}" if isinstance(before, (int, float)) else "-"
    a = f"{after:.3f}" if isinstance(after, (int, float)) else "-"
    return f"  {label}: before {b}, after {a}"


def _build_report_sections(request, include_bias_correction, include_water_levels):
    """
    Build report sections from the user's actual last analysis results (stashed
    in their session by the bias-correction/water-level views). There's no
    persistent analysis-history model, so this reflects the current session
    only — an honest "run it first" message when nothing's been run yet,
    rather than fabricated placeholder content.
    """
    sections = []

    if include_bias_correction:
        result = request.session.get('last_bias_result')
        if result:
            lines = [
                f"Bias Correction Analysis — method: {result['method']} "
                f"(variable: {result['variable']})"
            ]
            before = result['metrics_before']
            after = result['metrics_after']
            for key in ['RMSE', 'MAE', 'Bias', 'Correlation', 'NSE', 'KGE']:
                lines.append(_metric_line(key, before.get(key), after.get(key)))
            sections.append(lines)
        else:
            sections.append([
                "Bias Correction Analysis — no bias correction has been run yet "
                "this session. Run one from the Bias Correction tool first."
            ])

    if include_water_levels:
        result = request.session.get('last_levels_result')
        if result:
            sections.append([
                f"Water Level Prediction — {result['start_date']} to {result['end_date']}",
                f"  Minimum: {result['min_level']:.2f} m",
                f"  Maximum: {result['max_level']:.2f} m",
                f"  Average: {result['mean_level']:.2f} m",
            ])
        else:
            sections.append([
                "Water Level Prediction — no prediction has been run yet this "
                "session. Run one from the Water Level Prediction tool first."
            ])

    return sections


# __________________________________________________________________________________________________________REPORTS_VIEW____
@login_required
def reports(request):
    if request.user.is_authenticated:
        template_name = 'base_usr.html'
    else:
        template_name = 'base_all.html'

    if request.method == 'POST':
        include_bias_correction = 'bias_correction' in request.POST
        include_water_levels = 'water_levels' in request.POST
        export_format = request.POST.get('export_format', 'pdf')

        sections = _build_report_sections(request, include_bias_correction, include_water_levels)

        # Export as Word document
        if export_format == 'word':
            document = Document()
            document.add_heading('Analysis Report', level=1)
            for lines in sections:
                document.add_heading(lines[0], level=2)
                for line in lines[1:]:
                    document.add_paragraph(line)
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="analysis_report.docx"'
            return response

        # Export as PDF
        elif export_format == 'pdf':
            buffer = BytesIO()
            pdf = canvas.Canvas(buffer)
            pdf.setFont("Helvetica-Bold", 16)
            pdf.drawString(72, 800, "Analysis Report")
            y = 765
            for lines in sections:
                pdf.setFont("Helvetica-Bold", 12)
                pdf.drawString(72, y, lines[0])
                y -= 20
                pdf.setFont("Helvetica", 11)
                for line in lines[1:]:
                    pdf.drawString(72, y, line)
                    y -= 16
                y -= 14
            pdf.save()
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="analysis_report.pdf"'
            return response

    return render(request, 'tools/reports.html', {'template_name': template_name})
